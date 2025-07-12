#!/usr/bin/env python3
"""
FIXED RAG Document Processor für RagFlow mit umfassendem Debugging
Behebt Projekt-Zuordnung und Upload-Probleme
"""

import os
import uuid
import json
from pathlib import Path
from typing import Optional, List, Dict, Any
import asyncio
from datetime import datetime

# Document processing libraries
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
    print("✅ PyPDF2 verfügbar")
except ImportError:
    PYPDF2_AVAILABLE = False
    print("⚠️ PyPDF2 nicht verfügbar - pip install PyPDF2")

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
    print("✅ pdfplumber verfügbar")
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    print("⚠️ pdfplumber nicht verfügbar - pip install pdfplumber")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
    print("✅ python-docx verfügbar")
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx nicht verfügbar - pip install python-docx")

class DebugLogger:
    """Debugging Logger für bessere Fehlerdiagnose"""
    
    def __init__(self, debug_file: str = "debug_document_processing.log"):
        self.debug_file = debug_file
        self.debug_data = []
    
    def log(self, level: str, message: str, data: Any = None):
        """Log debug information"""
        timestamp = datetime.utcnow().isoformat()
        log_entry = {
            "timestamp": timestamp,
            "level": level,
            "message": message,
            "data": data
        }
        
        self.debug_data.append(log_entry)
        
        # Console output mit Emojis
        emoji_map = {
            "INFO": "ℹ️",
            "SUCCESS": "✅", 
            "WARNING": "⚠️",
            "ERROR": "❌",
            "DEBUG": "🔍"
        }
        
        emoji = emoji_map.get(level, "📝")
        print(f"{emoji} [{level}] {message}")
        
        if data:
            print(f"    💾 Data: {data}")
    
    def save_debug_log(self):
        """Save debug log to file"""
        try:
            with open(self.debug_file, 'w', encoding='utf-8') as f:
                json.dump(self.debug_data, f, indent=2, ensure_ascii=False)
            print(f"📄 Debug log saved to: {self.debug_file}")
        except Exception as e:
            print(f"❌ Failed to save debug log: {e}")
    
    def get_debug_summary(self) -> Dict[str, Any]:
        """Get summary of debug session"""
        levels = {}
        for entry in self.debug_data:
            level = entry["level"]
            levels[level] = levels.get(level, 0) + 1
        
        return {
            "total_entries": len(self.debug_data),
            "levels": levels,
            "session_start": self.debug_data[0]["timestamp"] if self.debug_data else None,
            "session_end": self.debug_data[-1]["timestamp"] if self.debug_data else None
        }

class FixedDocumentProcessor:
    """FIXED Document Processor mit Debugging und korrekter Projekt-Zuordnung"""
    
    def __init__(self, debug_mode: bool = True):
        self.debug = DebugLogger() if debug_mode else None
        self.supported_types = {
            'pdf': self._extract_pdf_text,
            'docx': self._extract_docx_text,
            'txt': self._extract_txt_text,
            'md': self._extract_txt_text
        }
        
        if self.debug:
            self.debug.log("INFO", "FixedDocumentProcessor initialisiert", {
                "supported_types": list(self.supported_types.keys()),
                "pdf_available": PDFPLUMBER_AVAILABLE or PYPDF2_AVAILABLE,
                "docx_available": DOCX_AVAILABLE
            })
    
    async def process_document_with_project(self, file_path: str, file_type: str, 
                                          document_id: str, project_id: str,
                                          documents_db: dict) -> Dict[str, Any]:
        """
        FIXED: Verarbeitet Dokument UND ordnet es korrekt dem Projekt zu
        
        Args:
            file_path: Pfad zur Datei
            file_type: Dateityp
            document_id: Dokument-ID
            project_id: Projekt-ID (WICHTIG!)
            documents_db: Dokumenten-Datenbank
        """
        if self.debug:
            self.debug.log("INFO", f"Starte Dokumentverarbeitung mit Projekt-Zuordnung", {
                "file_path": file_path,
                "file_type": file_type,
                "document_id": document_id,
                "project_id": project_id
            })
        
        # 1. Prüfe ob Dokument in DB existiert
        if document_id not in documents_db:
            error_msg = f"Dokument-ID {document_id} nicht in Datenbank gefunden"
            if self.debug:
                self.debug.log("ERROR", error_msg, {"available_docs": list(documents_db.keys())})
            return {"success": False, "error": error_msg}
        
        doc = documents_db[document_id]
        
        # 2. FIXED: Stelle sicher dass project_ids existiert
        if "project_ids" not in doc:
            doc["project_ids"] = []
            if self.debug:
                self.debug.log("WARNING", "project_ids fehlte - wurde erstellt")
        
        # 3. FIXED: Füge Projekt-ID hinzu wenn nicht vorhanden
        if project_id not in doc["project_ids"]:
            doc["project_ids"].append(project_id)
            if self.debug:
                self.debug.log("SUCCESS", f"Projekt {project_id} zu Dokument hinzugefügt", {
                    "document_project_ids": doc["project_ids"]
                })
        
        # 4. Verarbeite das Dokument
        try:
            doc["processing_status"] = "processing"
            
            result = await self.process_document(file_path, file_type)
            
            if result["success"]:
                # Erfolgreiche Verarbeitung
                doc["processing_status"] = "completed"
                doc["extracted_text"] = result["text"]
                doc["text_chunks"] = result["chunks"]
                doc["text_metadata"] = result["metadata"]
                doc["processing_error"] = None
                doc["processed_at"] = datetime.utcnow().isoformat()
                
                if self.debug:
                    self.debug.log("SUCCESS", f"Dokument erfolgreich verarbeitet", {
                        "filename": doc.get("filename"),
                        "word_count": result["metadata"].get("word_count"),
                        "chunk_count": len(result["chunks"]),
                        "project_ids": doc["project_ids"]
                    })
                
                return {
                    "success": True,
                    "document_id": document_id,
                    "project_id": project_id,
                    "text_preview": result["text"][:200] + "..." if len(result["text"]) > 200 else result["text"],
                    "metadata": result["metadata"]
                }
            else:
                # Fehler bei Verarbeitung
                doc["processing_status"] = "failed"
                doc["processing_error"] = result["error"]
                doc["extracted_text"] = ""
                doc["text_chunks"] = []
                
                if self.debug:
                    self.debug.log("ERROR", f"Dokumentverarbeitung fehlgeschlagen", {
                        "filename": doc.get("filename"),
                        "error": result["error"]
                    })
                
                return {"success": False, "error": result["error"]}
                
        except Exception as e:
            error_msg = f"Unerwarteter Fehler bei Verarbeitung: {str(e)}"
            doc["processing_status"] = "failed"
            doc["processing_error"] = error_msg
            
            if self.debug:
                self.debug.log("ERROR", error_msg, {"exception_type": type(e).__name__})
            
            return {"success": False, "error": error_msg}
    
    async def process_document(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """Basis-Dokumentverarbeitung (bestehende Logik)"""
        try:
            if self.debug:
                self.debug.log("INFO", f"Starte Basis-Verarbeitung", {
                    "file_path": file_path,
                    "file_type": file_type,
                    "file_exists": os.path.exists(file_path)
                })
            
            if not os.path.exists(file_path):
                return {
                    "success": False,
                    "error": f"Datei nicht gefunden: {file_path}",
                    "text": "",
                    "chunks": [],
                    "metadata": {}
                }
            
            if file_type not in self.supported_types:
                return {
                    "success": False,
                    "error": f"Dateityp '{file_type}' wird nicht unterstützt",
                    "text": "",
                    "chunks": [],
                    "metadata": {}
                }
            
            # Text extrahieren
            extractor = self.supported_types[file_type]
            text = await extractor(file_path)
            
            if not text or len(text.strip()) < 10:
                return {
                    "success": False,
                    "error": "Kein Text extrahiert oder Text zu kurz",
                    "text": "",
                    "chunks": [],
                    "metadata": {}
                }
            
            # Text in Chunks aufteilen
            chunks = self._chunk_text(text)
            
            # Metadaten erstellen
            metadata = {
                "word_count": len(text.split()),
                "char_count": len(text),
                "chunk_count": len(chunks),
                "extraction_method": file_type,
                "processed_at": datetime.utcnow().isoformat()
            }
            
            if self.debug:
                self.debug.log("SUCCESS", "Basis-Verarbeitung erfolgreich", metadata)
            
            return {
                "success": True,
                "error": None,
                "text": text,
                "chunks": chunks,
                "metadata": metadata
            }
            
        except Exception as e:
            error_msg = f"Fehler bei der Verarbeitung: {str(e)}"
            if self.debug:
                self.debug.log("ERROR", error_msg, {"exception_type": type(e).__name__})
            
            return {
                "success": False,
                "error": error_msg,
                "text": "",
                "chunks": [],
                "metadata": {}
            }
    
    async def _extract_pdf_text(self, file_path: str) -> str:
        """Extrahiert Text aus PDF-Dateien mit Debugging"""
        if self.debug:
            self.debug.log("INFO", "PDF-Textextraktion gestartet", {
                "pdfplumber_available": PDFPLUMBER_AVAILABLE,
                "pypdf2_available": PYPDF2_AVAILABLE
            })
        
        try:
            # Versuch 1: pdfplumber
            if PDFPLUMBER_AVAILABLE:
                try:
                    import pdfplumber
                    with pdfplumber.open(file_path) as pdf:
                        text_parts = []
                        for page_num, page in enumerate(pdf.pages):
                            try:
                                page_text = page.extract_text()
                                if page_text and page_text.strip():
                                    cleaned_text = self._clean_extracted_text(page_text)
                                    if cleaned_text:
                                        text_parts.append(f"[Seite {page_num + 1}]\n{cleaned_text}")
                            except Exception as page_error:
                                if self.debug:
                                    self.debug.log("WARNING", f"Fehler bei Seite {page_num + 1}", str(page_error))
                                continue
                        
                        full_text = "\n\n".join(text_parts)
                        if full_text.strip():
                            if self.debug:
                                self.debug.log("SUCCESS", f"pdfplumber erfolgreich", {
                                    "pages_extracted": len(text_parts),
                                    "total_chars": len(full_text)
                                })
                            return full_text
                except Exception as e:
                    if self.debug:
                        self.debug.log("WARNING", f"pdfplumber Fehler: {e}")
            
            # Versuch 2: PyPDF2
            if PYPDF2_AVAILABLE:
                try:
                    import PyPDF2
                    with open(file_path, 'rb') as file:
                        reader = PyPDF2.PdfReader(file)
                        text_parts = []
                        
                        for page_num, page in enumerate(reader.pages):
                            try:
                                page_text = page.extract_text()
                                if page_text and page_text.strip():
                                    cleaned_text = self._clean_extracted_text(page_text)
                                    if cleaned_text:
                                        text_parts.append(f"[Seite {page_num + 1}]\n{cleaned_text}")
                            except Exception as page_error:
                                if self.debug:
                                    self.debug.log("WARNING", f"PyPDF2 Fehler bei Seite {page_num + 1}", str(page_error))
                                continue
                        
                        full_text = "\n\n".join(text_parts)
                        if full_text.strip():
                            if self.debug:
                                self.debug.log("SUCCESS", f"PyPDF2 erfolgreich", {
                                    "pages_extracted": len(text_parts),
                                    "total_chars": len(full_text)
                                })
                            return full_text
                except Exception as e:
                    if self.debug:
                        self.debug.log("WARNING", f"PyPDF2 Fehler: {e}")
            
            error_msg = "Keine funktionsfähige PDF-Bibliothek gefunden"
            if self.debug:
                self.debug.log("ERROR", error_msg)
            return ""
            
        except Exception as e:
            if self.debug:
                self.debug.log("ERROR", f"PDF Extraction Fehler: {e}")
            return ""
    
    async def _extract_docx_text(self, file_path: str) -> str:
        """Extrahiert Text aus DOCX-Dateien mit Debugging"""
        if not DOCX_AVAILABLE:
            error_msg = "DOCX-Verarbeitung nicht verfügbar"
            if self.debug:
                self.debug.log("ERROR", error_msg)
            return ""
            
        try:
            from docx import Document
            doc = Document(file_path)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extrahiere Text aus Tabellen
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    text_parts.append("\n[Tabelle]\n" + "\n".join(table_text))
            
            full_text = "\n\n".join(text_parts)
            if self.debug:
                self.debug.log("SUCCESS", f"DOCX Text extrahiert", {
                    "paragraphs": len([p for p in doc.paragraphs if p.text.strip()]),
                    "tables": len(doc.tables),
                    "total_chars": len(full_text)
                })
            
            return full_text
            
        except Exception as e:
            if self.debug:
                self.debug.log("ERROR", f"DOCX Extraction Fehler: {e}")
            return ""
    
    async def _extract_txt_text(self, file_path: str) -> str:
        """Extrahiert Text aus TXT/MD-Dateien mit Debugging"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
                if self.debug:
                    self.debug.log("SUCCESS", f"TXT Text extrahiert", {
                        "encoding": "utf-8",
                        "chars": len(text)
                    })
                return text
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                    if self.debug:
                        self.debug.log("SUCCESS", f"TXT Text extrahiert", {
                            "encoding": "latin-1",
                            "chars": len(text)
                        })
                    return text
            except Exception as e:
                if self.debug:
                    self.debug.log("ERROR", f"TXT Extraction Fehler: {e}")
                return ""
        except Exception as e:
            if self.debug:
                self.debug.log("ERROR", f"TXT Extraction Fehler: {e}")
            return ""
    
    def _clean_extracted_text(self, text: str) -> str:
        """Säubert extrahierten Text"""
        if not text:
            return ""
        
        import re
        text = re.sub(r' +', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        return text.strip()
    
    def _chunk_text(self, text: str, max_chunk_size: int = 1000, overlap: int = 100) -> List[Dict[str, Any]]:
        """Teilt Text in Chunks mit Debugging"""
        if len(text) <= max_chunk_size:
            chunk = {
                "id": str(uuid.uuid4()),
                "text": text,
                "start_pos": 0,
                "end_pos": len(text),
                "chunk_index": 0
            }
            if self.debug:
                self.debug.log("INFO", "Single chunk created", {"size": len(text)})
            return [chunk]
        
        chunks = []
        start = 0
        chunk_index = 0
        
        while start < len(text):
            end = min(start + max_chunk_size, len(text))
            
            # Versuche an Satzende zu teilen
            if end < len(text):
                last_sentence = text.rfind('.', start, end)
                if last_sentence > start + max_chunk_size * 0.5:
                    end = last_sentence + 1
                else:
                    last_space = text.rfind(' ', start, end)
                    if last_space > start + max_chunk_size * 0.5:
                        end = last_space
            
            chunk_text = text[start:end].strip()
            if chunk_text:
                chunks.append({
                    "id": str(uuid.uuid4()),
                    "text": chunk_text,
                    "start_pos": start,
                    "end_pos": end,
                    "chunk_index": chunk_index
                })
                chunk_index += 1
            
            start = max(start + 1, end - overlap)
        
        if self.debug:
            self.debug.log("INFO", f"Text chunking completed", {
                "total_chunks": len(chunks),
                "average_size": sum(len(c["text"]) for c in chunks) // len(chunks) if chunks else 0
            })
        
        return chunks

class FixedRAGChatEnhancer:
    """FIXED RAG Chat Enhancer mit besserer Debugging"""
    
    def __init__(self, documents_db: dict, debug_mode: bool = True):
        self.documents_db = documents_db
        self.debug = DebugLogger() if debug_mode else None
    
    def find_relevant_content(self, query: str, project_id: str, max_chunks: int = 3) -> List[Dict[str, Any]]:
        """FIXED: Findet relevante Inhalte mit detailliertem Debugging"""
        
        if self.debug:
            self.debug.log("INFO", f"Starte intelligente Suche", {
                "query": query,
                "project_id": project_id,
                "max_chunks": max_chunks
            })
        
        # 1. Finde alle Dokumente
        all_docs = list(self.documents_db.values())
        if self.debug:
            self.debug.log("DEBUG", f"Alle Dokumente in DB", {
                "total_documents": len(all_docs),
                "document_ids": [doc.get("id", "no_id") for doc in all_docs]
            })
        
        # 2. Filtere nach Projekt
        project_docs = [
            doc for doc in all_docs
            if project_id in doc.get("project_ids", [])
        ]
        
        if self.debug:
            self.debug.log("DEBUG", f"Dokumente im Projekt gefiltert", {
                "project_documents": len(project_docs),
                "project_doc_names": [doc.get("filename", "no_name") for doc in project_docs]
            })
        
        # 3. Filtere nach Verarbeitungsstatus
        processed_docs = [
            doc for doc in project_docs
            if doc.get("processing_status") == "completed" and doc.get("extracted_text")
        ]
        
        if self.debug:
            self.debug.log("DEBUG", f"Verarbeitete Dokumente gefiltert", {
                "processed_documents": len(processed_docs),
                "processed_doc_details": [
                    {
                        "filename": doc.get("filename"),
                        "status": doc.get("processing_status"),
                        "has_text": bool(doc.get("extracted_text")),
                        "text_length": len(doc.get("extracted_text", ""))
                    }
                    for doc in processed_docs
                ]
            })
        
        if not processed_docs:
            if self.debug:
                self.debug.log("WARNING", "Keine verarbeiteten Dokumente gefunden", {
                    "total_docs_in_project": len(project_docs),
                    "project_doc_statuses": [
                        {
                            "filename": doc.get("filename"),
                            "status": doc.get("processing_status"),
                            "has_project_ids": bool(doc.get("project_ids")),
                            "project_ids": doc.get("project_ids", [])
                        }
                        for doc in project_docs
                    ]
                })
            return []
        
        # 4. Durchsuche Inhalte
        relevant_content = []
        query_lower = query.lower()
        query_words = query_lower.split()
        
        for doc in processed_docs:
            if self.debug:
                self.debug.log("DEBUG", f"Durchsuche Dokument", {
                    "filename": doc.get("filename"),
                    "text_length": len(doc.get("extracted_text", ""))
                })
            
            full_text = doc.get("extracted_text", "")
            text_lower = full_text.lower()
            
            # Keyword-Suche
            total_matches = 0
            for word in query_words:
                matches = text_lower.count(word)
                total_matches += matches
            
            if total_matches > 0:
                # Finde Kontexte um Keywords
                positions = []
                for word in query_words:
                    start = 0
                    while True:
                        pos = text_lower.find(word, start)
                        if pos == -1:
                            break
                        positions.append(pos)
                        start = pos + 1
                
                # Extrahiere Kontexte
                for pos in positions[:max_chunks]:
                    context_start = max(0, pos - 200)
                    context_end = min(len(full_text), pos + 200)
                    context = full_text[context_start:context_end]
                    
                    relevant_content.append({
                        "text": context,
                        "source": doc.get("filename", "Unknown"),
                        "document_id": doc.get("id"),
                        "score": total_matches / len(query_words),
                        "match_count": total_matches
                    })
        
        # Sortiere und limitiere
        relevant_content.sort(key=lambda x: x["score"], reverse=True)
        result = relevant_content[:max_chunks]
        
        if self.debug:
            self.debug.log("SUCCESS", f"Relevante Inhalte gefunden", {
                "total_found": len(relevant_content),
                "returned": len(result),
                "result_sources": [item["source"] for item in result]
            })
        
        return result

# Test und Debug Funktionen
async def debug_document_processing(file_path: str, project_id: str, documents_db: dict):
    """Debug-Funktion zum Testen der Dokumentverarbeitung"""
    
    print("🧪 === DEBUG DOCUMENT PROCESSING ===")
    
    processor = FixedDocumentProcessor(debug_mode=True)
    
    # Erstelle Test-Dokument in DB wenn nicht vorhanden
    document_id = str(uuid.uuid4())
    file_type = Path(file_path).suffix.lower().lstrip('.')
    
    documents_db[document_id] = {
        "id": document_id,
        "filename": Path(file_path).name,
        "file_type": file_type,
        "file_path": file_path,
        "uploaded_at": datetime.utcnow().isoformat(),
        "processing_status": "pending"
    }
    
    print(f"📄 Test-Dokument erstellt: {document_id}")
    
    # Verarbeite Dokument
    result = await processor.process_document_with_project(
        file_path, file_type, document_id, project_id, documents_db
    )
    
    print("\n📊 === VERARBEITUNGS-ERGEBNIS ===")
    print(f"Erfolg: {result.get('success')}")
    if result.get('success'):
        print(f"Text-Vorschau: {result.get('text_preview', 'Keine Vorschau')}")
        print(f"Metadaten: {result.get('metadata', {})}")
    else:
        print(f"Fehler: {result.get('error')}")
    
    print("\n🔍 === DOKUMENT STATUS IN DB ===")
    doc = documents_db[document_id]
    print(f"Status: {doc.get('processing_status')}")
    print(f"Projekt-IDs: {doc.get('project_ids', [])}")
    print(f"Hat Text: {bool(doc.get('extracted_text'))}")
    print(f"Text-Länge: {len(doc.get('extracted_text', ''))}")
    
    # Teste RAG-Suche
    print("\n🧠 === RAG-SUCHE TEST ===")
    rag_enhancer = FixedRAGChatEnhancer(documents_db, debug_mode=True)
    
    test_queries = ["test", "inhalt", doc.get("filename", "").split(".")[0]]
    
    for query in test_queries:
        print(f"\n🔍 Teste Query: '{query}'")
        results = rag_enhancer.find_relevant_content(query, project_id)
        print(f"Gefunden: {len(results)} Ergebnisse")
        for i, result in enumerate(results):
            print(f"  {i+1}. {result['source']} (Score: {result['score']:.2f})")
    
    # Speichere Debug-Log
    processor.debug.save_debug_log()
    rag_enhancer.debug.save_debug_log()
    
    print("\n✅ Debug-Session abgeschlossen!")
    return result

if __name__ == "__main__":
    # Beispiel-Nutzung
    print("🚀 Fixed Document Processor mit Debugging gestartet")
    
    # Test mit einer Beispiel-Datei
    test_file = "test_document.txt"
    test_project = "af4aqupco"
    test_db = {}
    
    # Erstelle Test-Datei falls nicht vorhanden
    if not os.path.exists(test_file):
        with open(test_file, 'w', encoding='utf-8') as f:
            f.write("Dies ist ein Test-Dokument für das Debugging.\n\nEs enthält mehrere Absätze und Informationen.")
        print(f"📝 Test-Datei erstellt: {test_file}")
    
    # Führe Debug-Test aus
    asyncio.run(debug_document_processing(test_file, test_project, test_db))