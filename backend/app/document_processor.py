# backend/app/document_processor.py - Modern ChromaDB-integrated Document Processor
"""
Modern Document Processor for RagFlow
Handles advanced document processing and integrates with ChromaDB
"""

import io
import os
import mimetypes
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple, Union
from datetime import datetime

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import chardet
    CHARDET_AVAILABLE = True
except ImportError:
    CHARDET_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

from rich.console import Console
from .database import get_db_manager
from .config import settings

console = Console()


class DocumentProcessor:
    """Advanced document processor with ChromaDB integration"""
    
    def __init__(self):
        self.db = get_db_manager()
        self.supported_formats = self._get_supported_formats()
        
    def _get_supported_formats(self) -> Dict[str, Dict[str, Any]]:
        """Get supported document formats"""
        formats = {
            ".txt": {
                "name": "Plain Text",
                "processor": self._process_text,
                "available": True,
                "mime_types": ["text/plain"]
            },
            ".md": {
                "name": "Markdown",
                "processor": self._process_markdown,
                "available": MARKDOWN_AVAILABLE,
                "mime_types": ["text/markdown", "text/x-markdown"]
            },
            ".html": {
                "name": "HTML",
                "processor": self._process_html,
                "available": BS4_AVAILABLE,
                "mime_types": ["text/html"]
            },
            ".pdf": {
                "name": "PDF Document",
                "processor": self._process_pdf,
                "available": PDF_AVAILABLE,
                "mime_types": ["application/pdf"]
            },
            ".docx": {
                "name": "Word Document",
                "processor": self._process_docx,
                "available": DOCX_AVAILABLE,
                "mime_types": ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]
            },
            ".doc": {
                "name": "Word Document (Legacy)",
                "processor": self._process_doc_legacy,
                "available": False,  # Requires additional libraries
                "mime_types": ["application/msword"]
            }
        }
        
        return formats
    
    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions"""
        return [ext for ext, info in self.supported_formats.items() if info["available"]]
    
    def is_supported(self, file_path: Union[str, Path]) -> bool:
        """Check if file format is supported"""
        extension = Path(file_path).suffix.lower()
        return extension in self.get_supported_extensions()
    
    def detect_file_type(self, file_path: Union[str, Path], content: bytes = None) -> Dict[str, Any]:
        """Detect file type and encoding"""
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        # Get MIME type
        mime_type, _ = mimetypes.guess_type(str(file_path))
        
        # Detect encoding for text files
        encoding = "utf-8"
        confidence = 1.0
        
        if content and CHARDET_AVAILABLE:
            try:
                detected = chardet.detect(content)
                if detected["encoding"]:
                    encoding = detected["encoding"]
                    confidence = detected["confidence"]
            except Exception:
                pass
        
        # Get format info
        format_info = self.supported_formats.get(extension, {
            "name": "Unknown",
            "available": False
        })
        
        return {
            "extension": extension,
            "mime_type": mime_type,
            "encoding": encoding,
            "encoding_confidence": confidence,
            "format_name": format_info["name"],
            "supported": format_info.get("available", False),
            "processor": format_info.get("processor")
        }
    
    async def process_document(
        self, 
        file_path: Union[str, Path], 
        project_ids: List[str] = None,
        metadata: Dict[str, Any] = None
    ) -> Dict[str, Any]:
        """Process document and store in ChromaDB"""
        
        file_path = Path(file_path)
        
        try:
            # Read file content
            with open(file_path, 'rb') as f:
                file_content = f.read()
            
            # Detect file type
            file_info = self.detect_file_type(file_path, file_content)
            
            if not file_info["supported"]:
                raise ValueError(f"Unsupported file format: {file_info['extension']}")
            
            console.print(f"📄 Processing {file_info['format_name']}: {file_path.name}")
            
            # Extract text content
            processor = file_info["processor"]
            text_content = await processor(file_content, file_info)
            
            # Prepare metadata
            doc_metadata = {
                "original_filename": file_path.name,
                "file_extension": file_info["extension"],
                "mime_type": file_info["mime_type"],
                "encoding": file_info["encoding"],
                "encoding_confidence": file_info["encoding_confidence"],
                "file_size_bytes": len(file_content),
                "processed_at": datetime.utcnow().isoformat(),
                "processor_version": "2.0",
                "text_length": len(text_content),
                "estimated_reading_time": self._estimate_reading_time(text_content),
                **(metadata or {})
            }
            
            # Create document in ChromaDB
            document = self.db.create_document(
                filename=file_path.name,
                file_type=file_info["extension"],
                file_size=len(file_content),
                file_path=str(file_path.absolute()),
                content=text_content,
                project_ids=project_ids or []
            )
            
            # Update document metadata
            document.metadata.update(doc_metadata)
            self.db._update_document_metadata(document.id, document)
            
            console.print(f"✅ Document processed: {document.id}")
            
            return {
                "document_id": document.id,
                "filename": file_path.name,
                "file_type": file_info["extension"],
                "file_size": len(file_content),
                "text_length": len(text_content),
                "processing_status": "completed",
                "metadata": doc_metadata,
                "chunks_created": self._estimate_chunks(text_content),
                "project_ids": project_ids or []
            }
            
        except Exception as e:
            console.print(f"❌ Error processing {file_path.name}: {e}")
            raise Exception(f"Document processing failed: {str(e)}")
    
    async def _process_text(self, content: bytes, file_info: Dict[str, Any]) -> str:
        """Process plain text file"""
        try:
            encoding = file_info.get("encoding", "utf-8")
            text = content.decode(encoding, errors="replace")
            return self._clean_text(text)
        except Exception as e:
            raise Exception(f"Text processing failed: {e}")
    
    async def _process_markdown(self, content: bytes, file_info: Dict[str, Any]) -> str:
        """Process Markdown file"""
        try:
            encoding = file_info.get("encoding", "utf-8")
            md_text = content.decode(encoding, errors="replace")
            
            if MARKDOWN_AVAILABLE:
                # Convert Markdown to HTML, then extract text
                html = markdown.markdown(md_text)
                if BS4_AVAILABLE:
                    soup = BeautifulSoup(html, 'html.parser')
                    text = soup.get_text()
                else:
                    # Simple HTML tag removal
                    import re
                    text = re.sub(r'<[^>]+>', '', html)
            else:
                # Use raw markdown
                text = md_text
            
            return self._clean_text(text)
        except Exception as e:
            raise Exception(f"Markdown processing failed: {e}")
    
    async def _process_html(self, content: bytes, file_info: Dict[str, Any]) -> str:
        """Process HTML file"""
        try:
            encoding = file_info.get("encoding", "utf-8")
            html_content = content.decode(encoding, errors="replace")
            
            if BS4_AVAILABLE:
                soup = BeautifulSoup(html_content, 'html.parser')
                
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                
                # Extract text
                text = soup.get_text()
            else:
                # Simple HTML tag removal
                import re
                text = re.sub(r'<[^>]+>', '', html_content)
            
            return self._clean_text(text)
        except Exception as e:
            raise Exception(f"HTML processing failed: {e}")
    
    async def _process_pdf(self, content: bytes, file_info: Dict[str, Any]) -> str:
        """Process PDF file"""
        if not PDF_AVAILABLE:
            raise Exception("PDF processing requires PyPDF2. Install with: pip install PyPDF2")
        
        try:
            pdf_file = io.BytesIO(content)
            reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
                except Exception as e:
                    console.print(f"⚠️ Warning: Could not extract text from page {page_num + 1}: {e}")
                    text_parts.append(f"--- Page {page_num + 1} ---\n[Text extraction failed]")
            
            if not text_parts:
                raise Exception("No text could be extracted from PDF")
            
            full_text = "\n\n".join(text_parts)
            return self._clean_text(full_text)
            
        except Exception as e:
            raise Exception(f"PDF processing failed: {e}")
    
    async def _process_docx(self, content: bytes, file_info: Dict[str, Any]) -> str:
        """Process DOCX file"""
        if not DOCX_AVAILABLE:
            raise Exception("DOCX processing requires python-docx. Install with: pip install python-docx")
        
        try:
            docx_file = io.BytesIO(content)
            doc = DocxDocument(docx_file)
            
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_text.append(row_text)
                
                if table_text:
                    text_parts.append("\n".join(table_text))
            
            if not text_parts:
                raise Exception("No text could be extracted from DOCX")
            
            full_text = "\n\n".join(text_parts)
            return self._clean_text(full_text)
            
        except Exception as e:
            raise Exception(f"DOCX processing failed: {e}")
    
    async def _process_doc_legacy(self, content: bytes, file_info: Dict[str, Any]) -> str:
        """Process legacy DOC file (placeholder)"""
        raise Exception("Legacy DOC format not supported. Please convert to DOCX or use a specialized library.")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        if not text:
            return ""
        
        # Normalize whitespace
        import re
        
        # Replace multiple whitespace with single space
        text = re.sub(r'\s+', ' ', text)
        
        # Replace multiple newlines with double newline
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Remove excessive spaces around newlines
        text = re.sub(r' *\n *', '\n', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def _estimate_reading_time(self, text: str) -> int:
        """Estimate reading time in minutes (average 250 words per minute)"""
        if not text:
            return 0
        
        word_count = len(text.split())
        reading_time = max(1, round(word_count / 250))
        return reading_time
    
    def _estimate_chunks(self, text: str) -> int:
        """Estimate number of chunks this text will create"""
        if not text:
            return 0
        
        chunk_size = settings.chunk_size
        overlap = settings.chunk_overlap
        
        if len(text) <= chunk_size:
            return 1
        
        # Rough estimation
        effective_chunk_size = chunk_size - overlap
        estimated_chunks = max(1, (len(text) - overlap) // effective_chunk_size)
        
        return estimated_chunks
    
    def get_processing_stats(self) -> Dict[str, Any]:
        """Get document processing statistics"""
        stats = self.db.get_stats()
        
        return {
            "supported_formats": len(self.get_supported_extensions()),
            "available_formats": [
                {
                    "extension": ext,
                    "name": info["name"],
                    "available": info["available"]
                }
                for ext, info in self.supported_formats.items()
            ],
            "total_documents": stats.get("documents", {}).get("total", 0),
            "total_chunks": stats.get("documents", {}).get("chunks_total", 0),
            "capabilities": {
                "pdf_support": PDF_AVAILABLE,
                "docx_support": DOCX_AVAILABLE,
                "html_support": BS4_AVAILABLE,
                "markdown_support": MARKDOWN_AVAILABLE,
                "encoding_detection": CHARDET_AVAILABLE
            }
        }

    async def batch_process_directory(
        self, 
        directory_path: Union[str, Path],
        project_ids: List[str] = None,
        recursive: bool = True,
        file_pattern: str = "*"
    ) -> Dict[str, Any]:
        """Process all supported documents in a directory"""
        
        directory_path = Path(directory_path)
        
        if not directory_path.exists():
            raise ValueError(f"Directory does not exist: {directory_path}")
        
        # Find all files
        if recursive:
            files = list(directory_path.rglob(file_pattern))
        else:
            files = list(directory_path.glob(file_pattern))
        
        # Filter supported files
        supported_files = [f for f in files if f.is_file() and self.is_supported(f)]
        
        console.print(f"📁 Found {len(supported_files)} supported files in {directory_path}")
        
        results = {
            "total_files": len(files),
            "supported_files": len(supported_files),
            "processed": [],
            "failed": [],
            "skipped": []
        }
        
        for file_path in supported_files:
            try:
                console.print(f"📄 Processing: {file_path.name}")
                
                result = await self.process_document(
                    file_path=file_path,
                    project_ids=project_ids
                )
                
                results["processed"].append({
                    "file": str(file_path),
                    "document_id": result["document_id"],
                    "status": "success"
                })
                
            except Exception as e:
                console.print(f"❌ Failed to process {file_path.name}: {e}")
                results["failed"].append({
                    "file": str(file_path),
                    "error": str(e),
                    "status": "failed"
                })
        
        console.print(f"✅ Batch processing complete: {len(results['processed'])} processed, {len(results['failed'])} failed")
        
        return results


# Global processor instance
document_processor = DocumentProcessor()

# Development helper
if __name__ == "__main__":
    import json
    import asyncio
    
    print("📄 Document Processor Status:")
    processor = DocumentProcessor()
    stats = processor.get_processing_stats()
    print(json.dumps(stats, indent=2))
    
    print(f"\n✅ Supported Extensions: {processor.get_supported_extensions()}")
    
    # Test file type detection
    test_files = ["test.txt", "document.pdf", "presentation.docx", "readme.md"]
    print(f"\n🔍 File Type Detection Test:")
    for test_file in test_files:
        file_info = processor.detect_file_type(test_file)
        status = "✅" if file_info["supported"] else "❌"
        print(f"   {status} {test_file}: {file_info['format_name']}")